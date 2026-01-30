from django.shortcuts import render
from django.http import JsonResponse,HttpResponse
import re
import io
import os
import string
from django.core.files.storage import FileSystemStorage
from docx import Document
from PyPDF2 import PdfReader
from django.views.decorators.csrf import csrf_exempt
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
import html



try:
    import language_tool_python
    tool = language_tool_python.LanguageTool('en-US')
except:
    tool = None



def home(request):
    return render(request, "home.html")

def about(request):
    return render(request, "about.html")

def service(request):
    return render(request, "plans.html")

def contact(request):
    return render(request, "contact.html")

def pricing(request):
    return render(request, "pricing.html")

def guide(request):
    return render(request, "guide.html")

def upload_page(request):
    return render(request, "upload.html") 


def fix_punctuation(text):
    import re
    # Space after punctuation
    text = re.sub(r'([.!?,])([A-Za-z])', r'\1 \2', text)
    # Remove repeated punctuation
    text = re.sub(r'([!?.,])\1+', r'\1', text)
    # Capitalize after sentence end
    text = re.sub(r'([.!?]\s+)([a-z])', lambda m: m.group(1) + m.group(2).upper(), text)
    return text



# helper for checkbox flags
def is_checked(request, name):
    return str(request.POST.get(name, "")).lower() in ["on", "true", "1"]


# ---------------- GRAMMAR ---------------- #

def correct_grammar(text):
    if not tool:
        return text

    matches = tool.check(text)
    corrected = text
    offset = 0

    for m in matches:
        if not m.replacements:
            continue

        replacement = m.replacements[0]
        start = m.offset + offset
        end = start + m.error_length

        corrected = corrected[:start] + replacement + corrected[end:]
        offset += len(replacement) - m.error_length

    return corrected

def upload_page(request):
    return render(request, "upload.html")  # your existing template

@csrf_exempt
def upload_file(request):
    if request.method != "POST":
        return JsonResponse({"success": False, "error": "Invalid request method"})

    uploaded_file = request.FILES.get("file")
    if not uploaded_file:
        return JsonResponse({"success": False, "error": "No file uploaded"})

    filename = uploaded_file.name
    ext = os.path.splitext(filename)[1].lower()

    try:
        text = ""
        if ext == ".pdf":
            reader = PdfReader(uploaded_file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        elif ext == ".docx":
            doc = Document(uploaded_file)
            text = "\n".join([p.text for p in doc.paragraphs])

        elif ext == ".txt":
            text = uploaded_file.read().decode("utf-8")

        else:
            return JsonResponse({"success": False, "error": "Unsupported file format"})

        return JsonResponse({"success": True, "text": text})

    except Exception as e:
        return JsonResponse({"success": False, "error": str(e)})








# ---------------- MAIN VIEW ---------------- #
def download_pdf(request):
    if request.method != "POST":
        return HttpResponse(status=405)

    text = request.POST.get("text", "")

    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="edited_document.pdf"'

    doc = SimpleDocTemplate(
        response,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40,
    )

    styles = getSampleStyleSheet()
    style = styles["Normal"]
    style.fontSize = 11
    style.leading = 16
    style.wordWrap = "CJK"  # 🔥 prevents long text overflow

    story = []

    for line in text.split("\n"):
        safe_line = html.escape(line)
        story.append(Paragraph(safe_line if safe_line else "&nbsp;", style))

    doc.build(story)

    return response

def download_docx(request):
    text = request.POST.get("text", "")

    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = 'attachment; filename="edited_document.docx"'

    return response

def download_txt(request):
    text = request.POST.get("text", "")

    response = HttpResponse(text, content_type="text/plain; charset=utf-8")
    response["Content-Disposition"] = 'attachment; filename="edited_document.txt"'

    return response

def process_text(request):
    result = {}
    text = request.POST.get("text", "")

    # ---------- FIND & REPLACE (FIRST) ----------
    find_word = request.POST.get("find_word", "")
    replace_word = request.POST.get("replace_word", "")

    if find_word:
        text = text.replace(find_word, replace_word)

    # ---------- CLEANING & FORMATTING ----------

    # Remove Extra Spaces
    if is_checked(request, "remove_spaces"):
        text = re.sub(r'\s+', ' ', text).strip()

    # Remove Emojis
    if is_checked(request, "remove_emoji"):
        emoji_pattern = re.compile(
            "[" 
            u"\U0001F600-\U0001F64F"
            u"\U0001F300-\U0001F5FF"
            u"\U0001F680-\U0001F6FF"
            u"\U0001F1E0-\U0001F1FF"
            "]+", flags=re.UNICODE
        )
        text = emoji_pattern.sub(r'', text)

    # Remove Empty Lines
    if is_checked(request, "remove_empty"):
        lines = [line for line in text.splitlines() if line.strip() != ""]
        text = "\n".join(lines)

    # Normalize Text
    if is_checked(request, "normalize_text"):
        replacements = {
            "“": '"', "”": '"',
            "‘": "'", "’": "'",
            "–": "-", "—": "-",
            "…": "...",
        }
        for k, v in replacements.items():
            text = text.replace(k, v)

    # Auto Capitalize Sentences
    if is_checked(request, "auto_caps"):
        sentences = re.split('([.!?]\s*)', text)
        result_text = ""
        for i in range(0, len(sentences), 2):
            part = sentences[i].capitalize()
            result_text += part
            if i + 1 < len(sentences):
                result_text += sentences[i + 1]
        text = result_text

    # Fix Repeated Punctuation
    if is_checked(request, "fix_punct"):
        text = re.sub(r'([!?\.])\1+', r'\1', text)

    # Remove URLs
    if is_checked(request, "remove_urls"):
        text = re.sub(r'http\S+|www\S+', '', text)

    # Remove HTML
    if is_checked(request, "remove_html"):
        text = re.sub(r'<.*?>', '', text)
    
    if is_checked(request, "fix_punct"):
        text = fix_punctuation(text)

    # Remove Duplicate Lines
    if is_checked(request, "remove_duplicates"):
        seen = set()
        new_lines = []
        for line in text.splitlines():
            if line not in seen:
                seen.add(line)
                new_lines.append(line)
        text = "\n".join(new_lines)

    # Remove Punctuation
    if is_checked(request, "remove_punct"):
        text = text.translate(str.maketrans('', '', string.punctuation))

    # ---------- CASE CONVERSION ----------

    action = request.POST.get("action")

    if action == "upper":
        text = text.upper()

    elif action == "lower":
        text = text.lower()

    elif action == "title":
        text = text.title()

    elif action == "toggle":
        text = "".join(
            c.lower() if c.isupper() else c.upper()
            for c in text
        )

    elif action == "sentence":
        parts = re.split(r'([.!?]\s*)', text or "")
        new_text = []

        for i in range(0, len(parts), 2):
            sentence = parts[i].strip()
            if sentence:
                sentence = sentence[0].upper() + sentence[1:]
            new_text.append(sentence)

            if i + 1 < len(parts):
                new_text.append(parts[i + 1])

        text = "".join(new_text)

    elif action == "snake":
        # snake_case
        words = re.findall(r'\b\w+\b', text.lower())
        text = "_".join(words)

    elif action == "kebab":
        # kebab-case
        words = re.findall(r'\b\w+\b', text.lower())
        text = "-".join(words)

    elif action == "camel":
        # camelCase
        words = re.findall(r'\b\w+\b', text.lower())
        if words:
            text = words[0] + "".join(w.capitalize() for w in words[1:])
        else:
            text = ""

    elif action == "pascal":
        # PascalCase
        words = re.findall(r'\b\w+\b', text.lower())
        text = "".join(w.capitalize() for w in words)

    elif action == "dot":
        # dot.case
        words = re.findall(r'\b\w+\b', text.lower())
        text = ".".join(words)

    # ---------- GRAMMAR (LAST STEP) ----------

    if is_checked(request, "grammar"):
        text = correct_grammar(text)

    # ---------- ANALYZER ----------

    words = len(text.split())
    characters = len(text)
    lines = len(text.splitlines())

    if is_checked(request, "count_words"):
        result["words"] = words

    if is_checked(request, "count_chars"):
        result["characters"] = characters

    if is_checked(request, "count_lines"):
        result["lines"] = lines

    if is_checked(request, "reading_time"):
        minutes = max(1, round(words / 200))
        result["reading_time"] = f"{minutes} min"

    # ---------- FINAL OUTPUT ----------

    result["text"] = text

    return JsonResponse({"text": text})

